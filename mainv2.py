import datetime
import pandas as pd
from pandas._libs.tslibs.parsing import DateParseError
import os
from docx import Document

#################### SETTINGS ########################
EXCEL_SHEET = 'No file selected.'
required_cols = [1,2,7,27,29]
save_path = os.getcwd()

document = Document()
# Set font
style = document.styles['Normal']
style.font.name = 'Calibri' # type: ignore
########################################################

def top_producer(producers: pd.DataFrame):
    print("Generating producers...")
    producers_dict = {}
    for i,j in producers.iterrows():
        if j["Consultant Name"] not in producers_dict:
            producers_dict[j["Consultant Name"]] = [j["Amount"], j["Branch"]]
        else:
            producers_dict[j["Consultant Name"]][0] += j["Amount"]
            # print("Added: ", j["Amount"], " to ", j["Consultant Name"])
            # assert producers_dict[j["Consultant Name"]][1] == j["Branch"]
    # print(producers_dict)
    sorted_producers = sorted(producers_dict.items(), key=lambda x: x[1][0], reverse=True)
    # print(sorted_producers)

    document.add_paragraph("Top Producers:")
    global total_amount
    total_amount = 0
    for producer in sorted_producers:
        total_amount += producer[1][0]
        currency = ('{:,}'.format(producer[1][0]))
        document.add_paragraph("{0}~ {1} - RM{2}".format(producer[1][1], producer[0], currency),\
                               style='List Number')
    document.add_paragraph()
    total_amount = ('{:,}\n'.format(total_amount))
    print("Producers generated!:", sorted_producers, sep="\n")

def top_case(cases: pd.DataFrame) -> None:
    print("Generating cases...")
    case_dict = {}
    for _,j in cases.iterrows():
        number_of_case = calculate_case(j["Amount"])
        if j["Consultant Name"] in case_dict:
            case_dict[j["Consultant Name"]][1] += number_of_case
        else:
            case_dict[j["Consultant Name"]] = [j["Branch"], number_of_case]
    # print(case_dict)
    sorted_cases = sorted(case_dict.items(), key=lambda x: x[1][1], reverse=True)
    # print(sorted_cases)
    
    document.add_paragraph("Top Cases:")
    global total_number_of_cases
    total_number_of_cases = 0
    for indiv_cases in sorted_cases:
        total_number_of_cases += indiv_cases[1][1]
        if indiv_cases[1][1] <= 1:
            document.add_paragraph("{0}~ {1} - {2} CASE".format(indiv_cases[1][0], indiv_cases[0], indiv_cases[1][1]),\
                                   style='List Number 2')
        else:
            document.add_paragraph("{0}~ {1} - {2} CASES".format(indiv_cases[1][0], indiv_cases[0], indiv_cases[1][1]),\
                                   style='List Number 2')
    document.add_paragraph()
    print("Cases generated!:", sorted_cases, sep="\n")
    
def branch(branches: pd.DataFrame) -> None:
    print("Generating branches...")
    branch_dict = {}
    for _,j in branches.iterrows():
        number_of_case = calculate_case(j["Amount"])
        if j["Branch"] not in branch_dict:
            branch_dict[j["Branch"]] = number_of_case
        else:
            branch_dict[j["Branch"]] += number_of_case
    # print(branch_dict)
    sorted_branches = sorted(branch_dict.items(), key=lambda x: x[1], reverse=True)
    # print(sorted_cases)

    document.add_paragraph("Top Branches:")
    for branch in sorted_branches:
        if branch[1] <= 1:
            document.add_paragraph("{0}~ {1} CASE".format(branch[0], branch[1]), style="List Number 3")
        else:
            document.add_paragraph("{0}~ {1} CASES".format(branch[0], branch[1]), style="List Number 3")
    document.add_paragraph()
    print("Branches generated!:", sorted_branches, sep="\n")

def calculate_case(amount):
    if amount in price_dict:
        return price_dict[amount]
    else:
        for multiplier in range(2,11):
            if amount%multiplier == 0:
                return multiplier
        return 1

price_dict: dict = {
    3299: 2,
    2804: 2,
    599: 1, 
    899: 1, 
    1499: 1, 
    1799: 1
}

def date(dates: pd.Series) -> list:
    print("Generating dates...")
    dates_formatted: pd.Series = pd.to_datetime(dates).dt.strftime("%d/%m/%Y")
    # print("Dates formatted:", dates_formatted)
    start_date = dates_formatted[0].split('/')
    start_date[1] = date_dict[start_date[1]]
    start_date = " ".join(start_date)
    end_date = dates_formatted[len(dates_formatted)-2].split('/')
    end_date[1] = date_dict[end_date[1]]
    end_date = " ".join(end_date)
    print("Dates generated!")
    print("Start date:", start_date)
    print("End date:", end_date)
    return [start_date, end_date]

date_dict = {
    "01": "Jan",
    "02" : "Feb",
    "03" : "Mar",
    "04" : "Apr",
    "05" : "May",
    "06" : "Jun",
    "07" : "Jul",
    "08" : "Aug",
    "09" : "Sep",
    "10" : "Oct",
    "11" : "Nov",
    "12" : "Dec"
}

def generate_report(custom_name: str):
    try:
        ########### TO WRITE ##############
        global report
        dt1 = datetime.datetime.now()

        print("Parsing excel sheet...:", EXCEL_SHEET)
        data = pd.read_excel(EXCEL_SHEET, usecols= required_cols)
        for i,j in data.iterrows():
            
            if pd.isnull(j["Status"]):
                raise ValueError("Empty cell in 'Status' column in excel sheet")
            if pd.isnull(j["Submission Date"]):
                raise ValueError("Empty cell in 'Submission Date' column in excel sheet")
            if pd.isnull(j["Amount"]):
                raise ValueError("Empty cell in 'Amount' column in excel sheet")
            if pd.isnull(j["Consultant Name"]):
                raise ValueError("Empty cell in 'Consultant Name' column in excel sheet")
            if pd.isnull(j["Branch"]):
                raise ValueError("Empty cell in 'Branch' column in excel sheet")
            
            if j["Status"] == "REJECTED":
                data = data.drop(i)
            data = data.replace(j["Branch"], j["Branch"].split(" ").pop())
        print("Parsing successful!", "Dataframe:", data, sep='\n')

        print("Generating submission dates...")
        date_output = date(data["Submission Date"])
        print("Submission dates generated!:", date_output)

        print("Generating filename...")
        if not custom_name:
            filename = "Production Report {0}".format(datetime.datetime.now().isoformat(timespec='seconds'))
            filename = filename.replace(":", '-')
        else:
            filename = "{0}".format(custom_name)

        filename = os.path.join(save_path, filename + ".docx")
        print("Filename generated!:", filename)

        print("Generating report...")
        para_header = document.add_paragraph("Month-to-date Production:\n{0} - {1}\nRM "\
                                             .format(date_output[0], date_output[1]))
        
        top_producer(data[["Branch","Consultant Name", "Amount"]])
        para_header.add_run(str(total_amount))
        top_case(data[["Branch", "Consultant Name", "Amount"]])
        branch(data[["Branch", "Amount"]])
        document.add_paragraph("TOTAL CASES: "+ str(total_number_of_cases)+ " CASES\n")

        document.add_paragraph("Your dedication to protecting and preserving family wealth distribution is truly \
commendable. You are not just selling products; you are safeguarding the future and \
ensuring the well-being of countless families. \n\nKeep up the fantastic work, and may \
your efforts continue to help secure the financial legacies of many more families.\
\n\nOnce again, congratulations on your well-deserved success!")
        
        document.save(filename)

    except FileNotFoundError as e:
        report = ("Error: File not found. Please check the provided path.\nInfo: {0}").format(e)
        print(report)
        return

    except DateParseError as e:
        report = ("Error: Date format not recognized. Please use DD-MM-YY.\nInfo: {0}").format(e)
        print(report)
        return

    except ValueError as e:
        report = ("Error: File format not recognized or cells could be empty. Please check file format or empty cells.\nInfo: {0}").format(e)
        print(report)
        return
    
    except Exception as ex:
        template = "An exception of type {0} occurred. Arguments:\n{1!r}"
        message = template.format(type(ex).__name__, ex.args)
        print(message)
    
    else:
        dt2 = datetime.datetime.now()
        td = (dt2 - dt1).total_seconds() *10**3
        td = str(round(td, 2))
        report = ("Report done in {0} ms.\nThe file is saved as: {1}".format(td, filename))
        print(report)
        return report


if __name__ == "__main__": 
    generate_report("")